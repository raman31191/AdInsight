import sys
import subprocess
import os

def install_and_import(package):
    try:
        import docx
    except ImportError:
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
        import docx

install_and_import('python-docx')
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    
    # Helper functions
    def add_h1(text):
        h = doc.add_heading(text, level=1)
        h.runs[0].font.size = Pt(22)
        h.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
    def add_h2(text):
        h = doc.add_heading(text, level=2)
        h.runs[0].font.size = Pt(16)
        h.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        
    def add_p(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing = 1.5
        
    def add_placeholder(title, instructions):
        doc.add_paragraph('\n')
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run(f"\n\n[ SCREENSHOT PLACEHOLDER ]\n{title}\n")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(200, 0, 0)
        
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"Action required by BA Intern: {instructions}\n(Insert image here to replace this block)\n\n")
        run2.italic = True
        run2.font.color.rgb = RGBColor(100, 100, 100)
        
        # Add massive spacing to simulate an image taking up the rest of the page
        for _ in range(25):
            doc.add_paragraph()
            
    # --- TITLE PAGE ---
    for _ in range(8): doc.add_paragraph()
    title = doc.add_paragraph('Edxellence')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(40)
    title.runs[0].bold = True
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('Meridian Analytics Platform\nOperations, Maintenance, and Strategic Impact Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(24)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    for _ in range(10): doc.add_paragraph()
    info = doc.add_paragraph('Prepared by: Business Analyst Intern\nCompany: Edxellence\nProject: Meridian Dashboard Maintenance\nDate: April 2026')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.runs[0].font.size = Pt(14)
    info.runs[0].bold = True
    
    doc.add_page_break()
    
    # --- TOC Placeholder ---
    add_h1("Table of Contents")
    add_p("1. Executive Summary & Edxellence Corporate Objectives")
    add_p("2. Introduction: The Business Analyst Role in Maintaining Meridian")
    add_p("3. Architectural Overview & Strategic Value")
    add_p("4. Core Dashboard Features & Business Motives")
    add_p("5. Real-Time Campaign Monitoring: Tactical Operations")
    add_p("6. Automated Anomaly Detection & Fraud Mitigation Strategies")
    add_p("7. SRE (Site Reliability Engineering) Observability & System Health")
    add_p("8. Supabase Backend Management & Data Integrity")
    add_p("9. Recent Feature Enhancements & Strategic Impact")
    add_p("10. Future Roadmap & Continuous Improvement Lifecycle")
    doc.add_page_break()
    
    # --- SECTION 1 ---
    add_h1("1. Executive Summary & Edxellence Corporate Objectives")
    add_p("At Edxellence, our primary mission is to leverage data-driven intelligence to maximize operational efficiency and marketing Return on Investment (ROI). The Meridian Campaign Intelligence Platform represents our flagship internal tool for paid media oversight. This report outlines the current state, maintenance protocols, and strategic enhancements implemented on the Meridian dashboard.")
    add_p("As a Business Analyst Intern, my core responsibility is to ensure this platform aligns with Edxellence’s business objectives: providing real-time visibility into ad spend, instantly detecting anomalies, and ensuring our pipeline architecture remains robust and reliable.")
    add_p("This comprehensive operational document details the functional components of the dashboard, the business motives driving each feature, our highly realistic campaign structures, and the rigorous backend maintenance performed via Supabase. It serves as both a status report to stakeholders and an operational runbook for the analytics team.")
    add_placeholder("Meridian Dashboard Overview", "Take a full-screen screenshot of the main Overview tab in your browser at http://localhost:3000 showing the KPIs, charts, and live feed.")
    doc.add_page_break()
    
    # --- SECTION 2 ---
    add_h1("2. Introduction: The Business Analyst Role in Maintaining Meridian")
    add_p("Maintaining an enterprise-grade analytics platform requires bridging the gap between raw data engineering and executive decision-making. As the assigned BA Intern at Edxellence, my maintenance tasks extend beyond technical monitoring; they involve interpreting the business value of the streaming data.")
    add_p("The Meridian platform is not a static tool; it is a live, dynamic pipeline. My role encompasses:")
    doc.add_paragraph("• Monitoring data integrity and ensuring the Python generator pipeline streams metrics accurately without latency.", style='List Bullet')
    doc.add_paragraph("• Evaluating the business logic of our automated alert systems to prevent alert fatigue among campaign managers.", style='List Bullet')
    doc.add_paragraph("• Translating technical backend (Supabase) configurations into tangible business outcomes.", style='List Bullet')
    doc.add_paragraph("• Overseeing the recent migration to 8 highly realistic, industry-benchmarked Google Ads campaign types.", style='List Bullet')
    add_placeholder("BA Workstation / Terminal View", "Take a screenshot showing the Python generator running in the terminal alongside the Vite React server terminal, illustrating the dual-system monitoring task.")
    doc.add_page_break()
    
    # --- SECTION 3 ---
    add_h1("3. Architectural Overview & Strategic Value")
    add_p("The architecture of Meridian was designed to prioritize real-time data delivery, a crucial requirement for Edxellence's fast-paced marketing operations. The system is divided into three critical tiers:")
    
    add_h2("Tier 1: Data Ingestion (Python Pipeline)")
    add_p("A robust Python script acts as our data generator, simulating real-world traffic patterns, including time-of-day peaks and day-to-day variances. It calculates complex metrics like ROAS natively before pushing to the cloud, ensuring standardization.")
    add_placeholder("Python Generator Code", "Take a screenshot of the generator.py file in VS Code showing the Time-of-Day multiplier logic.")
    doc.add_page_break()
    
    add_h2("Tier 2: Data Storage & Realtime Engine (Supabase)")
    add_p("Supabase serves as our PostgreSQL backend. We leverage its built-in WebSockets (Realtime) to broadcast database INSERTs directly to the client. This architectural choice saves Edxellence thousands of dollars in server costs by eliminating the need for a standalone WebSocket server.")
    add_placeholder("Supabase Project Dashboard", "Take a screenshot of the Supabase project dashboard (Table editor view) showing the 'ad_metrics' table populated with live data.")
    doc.add_page_break()
    
    add_h2("Tier 3: Client Visualization (React/Vite)")
    add_p("The frontend is built on React 18, utilizing Recharts for data visualization. It consumes the WebSocket stream to update the DOM without page refreshes, providing an 'always-live' command center.")
    add_placeholder("React Component Code", "Take a screenshot of App.jsx showing the Supabase Realtime subscription useEffect block.")
    doc.add_page_break()
    
    # --- SECTION 4 ---
    add_h1("4. Core Dashboard Features & Business Motives")
    add_p("Every UI element in Meridian was built with a specific Edxellence business motive. As the BA, I maintain these features to ensure they continue to deliver actionable insights.")
    
    features = [
        {"name": "Live Event Stream", "tech": "WebSocket subscription rendering a top-60 event list.", "motive": "Provides immediate qualitative feedback on traffic flow. Allows analysts to 'feel' the pulse of the campaigns and spot sudden halts in traffic visually."},
        {"name": "Top-Level KPI Cards", "tech": "Aggregated computations of spend, impressions, clicks, conversions, CTR, and ROAS.", "motive": "Executive summary. Allows Edxellence leadership to gauge daily performance at a single glance without digging into granular tables."},
        {"name": "Real-Time Area Chart", "tech": "Recharts AreaChart mapping the last 40 data points on a timeline.", "motive": "Identifies macro-trends. The visual correlation between spend spikes and conversion lifts helps analysts validate budget pacing."},
        {"name": "Campaign Spend Distribution (Bar Chart)", "tech": "Aggregates total spend per campaign into a comparative bar graph.", "motive": "Budget allocation monitoring. Ensures no single campaign is cannibalizing the daily budget maliciously."},
        {"name": "Active Alerts Feed", "tech": "Filtered view of the 'alerts' table where resolved=false.", "motive": "Risk mitigation. This is the most critical operational feature, routing anomalies directly to the human-in-the-loop for immediate action."}
    ]
    
    for f in features:
        add_h2(f['name'])
        add_p(f"Technical Implementation: {f['tech']}")
        add_p(f"Edxellence Business Motive: {f['motive']}")
        add_placeholder(f"Feature Snapshot: {f['name']}", f"Crop a screenshot of the dashboard focusing specifically on the {f['name']} component.")
        doc.add_page_break()

    # --- SECTION 5 ---
    add_h1("5. Real-Time Campaign Monitoring: Tactical Operations")
    add_p("A major enhancement I oversaw was migrating Meridian from generic test data to 8 highly realistic Google Ads campaign profiles. This section details the business logic behind each campaign I am currently maintaining for Edxellence.")
    
    campaigns = [
        {"name": "Search | Brand Terms", "intent": "High intent; capturing users explicitly searching for Edxellence.", "budget": "$27,000", "kpi": "CTR: 8-18%, CPC: $0.30 - $1.20"},
        {"name": "Search | Generic High-Intent", "intent": "Competitive market terms; capturing users looking for our service category.", "budget": "$84,000", "kpi": "CTR: 3-7%, CPC: $2.50 - $6.50"},
        {"name": "Search | Competitor Conquest", "intent": "Bidding on rival brand names to siphon traffic.", "budget": "$42,000", "kpi": "CTR: 1-3%, CPC: $4.50 - $10.00"},
        {"name": "Shopping | All Products", "intent": "Product listing ads for direct e-commerce conversions.", "budget": "$96,000", "kpi": "CTR: 0.8-2.5%, CPC: $0.40 - $1.80"},
        {"name": "Display | Prospecting", "intent": "Top-of-funnel brand awareness; massive reach.", "budget": "$48,000", "kpi": "CTR: 0.05-0.2%, CPC: $0.10 - $0.55"},
        {"name": "Display | Retargeting", "intent": "Re-engaging users who previously visited Edxellence domains.", "budget": "$28,500", "kpi": "CTR: 0.3-1.0%, CPC: $0.20 - $0.85"},
        {"name": "Performance Max | ROAS Target", "intent": "Google's AI-driven automated bidding across all networks.", "budget": "$135,000", "kpi": "CTR: 2-5%, CPC: $1.50 - $4.00"},
        {"name": "YouTube | Brand Awareness", "intent": "Video ads optimized for views, not direct clicks.", "budget": "$66,000", "kpi": "CTR: 0.03-0.15%, CPC: $0.03 - $0.18"}
    ]
    
    for c in campaigns:
        add_h2(c['name'])
        add_p(f"Business Intent: {c['intent']}")
        add_p(f"Allocated Budget: {c['budget']}")
        add_p(f"Benchmark KPIs: {c['kpi']}")
        add_p("Maintenance Focus: As an analyst, I monitor this campaign's daily spend velocity against its allocated budget. If the CPC spikes beyond the benchmark range, it triggers an investigation into keyword auction competitiveness.")
        add_placeholder(f"Campaign Analytics: {c['name']}", f"Go to the 'Campaigns' tab on the dashboard and take a screenshot of the specific table row for '{c['name']}'.")
        doc.add_page_break()

    # --- SECTION 6 ---
    add_h1("6. Automated Anomaly Detection & Fraud Mitigation Strategies")
    add_p("Edxellence loses thousands of dollars to click fraud and runaway budgets if not monitored. I maintain the Python pipeline's alert engine which executes four specific business rules:")
    add_p("1. Break-even ROAS Protection: Flags campaigns dropping below 2.0x ROAS.")
    add_p("2. CTR Baseline Drops: Flags campaigns where CTR drops 40% below historical minimums.")
    add_p("3. Budget Exhaustion: Alerts when 90% of a daily budget is consumed.")
    add_p("4. Traffic Spikes (Click Fraud): Detects sudden 10x spikes in impressions and clicks disjointed from normal conversion rates.")
    add_p("To prevent 'alert fatigue', I successfully implemented a 5-minute cooldown per campaign per alert type.")
    add_placeholder("Active Alerts Feed Snapshot", "Take a screenshot of the Active Alerts feed showing a RED (High Severity) or YELLOW (Medium Severity) alert generated by the anomaly engine.")
    doc.add_page_break()

    # --- SECTION 7 ---
    add_h1("7. SRE (Site Reliability Engineering) Observability & System Health")
    add_p("Analytics are useless if the data is stale. The 'System Health' tab in Meridian is my daily checklist for infrastructure reliability.")
    add_p("I monitor the 'pipeline_logs' table to ensure the Python generator successfully connects and inserts data. The frontend WebSocket connection status (Connected vs Disconnected) is critical. If offline, the business is flying blind.")
    add_placeholder("System Health Monitor Tab", "Take a full screenshot of the System Health (SRE) tab showing the Pipeline Run History, latest Run ID, Duration, and the LIVE status badge.")
    doc.add_page_break()

    # --- SECTION 8 ---
    add_h1("8. Supabase Backend Management & Data Integrity")
    add_p("As part of my maintenance duties, I actively manage the Supabase PostgreSQL instance. This involves:")
    add_p("- Managing Environment Variables: Ensuring frontend Anon keys and backend Service keys are secure and rotated when necessary.")
    add_p("- Executing SQL Scripts: Running seed data (e.g., the recent v2 campaign seed) securely via the Supabase SQL Editor.")
    add_p("- Monitoring Logical Replication: Ensuring the 'supabase_realtime' publication includes 'ad_metrics' and 'alerts' tables so the WebSockets function correctly.")
    add_p("- Database Consistency: Ensuring the computed 'ctr' column functions correctly at the database level rather than application level to prevent rounding errors.")
    add_placeholder("Supabase SQL Editor Execution", "Take a screenshot of the Supabase SQL Editor showing the 'seed_campaigns.sql' script execution and the 'Success' message.")
    doc.add_page_break()

    # --- SECTION 9 ---
    add_h1("9. Recent Feature Enhancements & Strategic Impact")
    add_p("Over the course of my maintenance, several critical features were deployed to elevate Meridian from a prototype to an enterprise-grade platform:")
    add_p("1. Complete Rebranding: Transformed 'AdInsight Live' to 'Meridian', implementing a professional, company-driven aesthetic suitable for Edxellence executives.")
    add_p("2. 8 Production Campaigns: Replaced dummy data with industry-benchmarked Google Ads profiles.")
    add_p("3. Infinite Loop Resolution: Refactored React 'useEffect' dependencies using 'useRef' to prevent frontend crashing during high-throughput WebSocket streams.")
    add_p("4. Schema Alignment: Ensured the frontend strictly reads from 'ad_metrics' and 'alerts' rather than deprecated testing tables.")
    add_placeholder("Meridian Branding Update", "Take a screenshot showing the top-left Navbar featuring the new 'Meridian' logo and 'Campaign Intelligence Platform' subtitle.")
    doc.add_page_break()

    # --- SECTION 10 ---
    add_h1("10. Future Roadmap & Continuous Improvement Lifecycle")
    add_p("As Edxellence scales, the Meridian platform must evolve. Based on my analysis as a BA intern, the following enhancements are proposed for the upcoming quarters:")
    add_p("1. Date Range Filtering: Implementing time-series queries for last 24h, 7d, and 30d analytics.")
    add_p("2. Alert Acknowledgement System: Allowing analysts to click 'Resolve' on an alert, mutating the Supabase 'resolved' boolean directly from the UI.")
    add_p("3. Machine Learning Forecasting: Integrating linear regression into the Python pipeline to predict end-of-day spend pacing.")
    add_p("4. Export to CSV: Adding a data portability feature for external reporting workflows.")
    add_placeholder("Code Architecture / IDE", "Take a screenshot of VS Code showing the full project directory structure (frontend, pipeline, database) to showcase the comprehensive full-stack nature of the project to your supervisor.")
    
    # --- CONCLUSION ---
    doc.add_page_break()
    add_h1("Conclusion")
    add_p("Maintaining the Meridian platform at Edxellence has demonstrated the critical intersection of data engineering, real-time analytics, and strategic business operations. Through rigorous monitoring, system upgrades, and alignment with corporate KPIs, the platform currently stands as a highly reliable, enterprise-grade tool driving our paid media success.")
    add_p("By leveraging industry-benchmarked data, robust anomaly detection, and real-time observability, the analytics team is now empowered to make proactive, budget-saving decisions.")
    add_p("Report Concluded.")

    doc.save('Edxellence_Meridian_Maintenance_Report.docx')

if __name__ == '__main__':
    create_report()
